from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "HUONG_DAN_SU_DUNG_ATG_PC_AUDIT.docx"
NAVY = "0B3A6E"; BLUE = "1F4E78"; PALE = "EAF2F8"; LIGHT = "F4F6F9"; GOLD = "D6A100"; RED = "B42318"; GREEN = "217346"

def font(run, size=11, bold=False, color="222222", italic=False):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr(); node = tc.first_child_found_in("w:tcMar")
    if node is None: node = OxmlElement("w:tcMar"); tc.append(node)
    for tag,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        x=OxmlElement("w:"+tag); x.set(qn("w:w"),str(val)); x.set(qn("w:type"),"dxa"); node.append(x)

def set_repeat(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); trPr.append(el)

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); p.paragraph_format.keep_with_next=True
    p.add_run(text); return p

def body(doc, text, bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.25
    if bold_prefix and text.startswith(bold_prefix): font(p.add_run(bold_prefix),bold=True,color=NAVY); text=text[len(bold_prefix):]
    font(p.add_run(text)); return p

def bullet(doc, text, level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2"); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.15
    font(p.add_run(text)); return p

CURRENT_NUM_ID = None
def start_numbering(doc):
    global CURRENT_NUM_ID
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num")); next_id = max([int(x.get(qn("w:numId"))) for x in nums] + [0]) + 1
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"), "7"); num.append(abstract)
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); num.append(override); numbering.append(num)
    CURRENT_NUM_ID = next_id

def number(doc, text):
    if CURRENT_NUM_ID is None: start_numbering(doc)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.15
    pPr=p._p.get_or_add_pPr(); numPr=OxmlElement("w:numPr"); ilvl=OxmlElement("w:ilvl"); ilvl.set(qn("w:val"),"0"); numId=OxmlElement("w:numId"); numId.set(qn("w:val"),str(CURRENT_NUM_ID)); numPr.append(ilvl); numPr.append(numId); pPr.append(numPr)
    font(p.add_run(text)); return p

def callout(doc, title, text, color=BLUE, fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.5)
    trPr=t.rows[0]._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); trPr.append(cant)
    c=t.cell(0,0); shade(c,fill); margins(c,120,160,120,160)
    p=c.paragraphs[0]; font(p.add_run(title+"\n"),bold=True,color=color); font(p.add_run(text),color="333333")
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    if widths:
        for i,w in enumerate(widths): t.columns[i].width=Inches(w)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,BLUE); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),bold=True,color="FFFFFF",size=10)
    set_repeat(t.rows[0])
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if ridx%2: shade(cells[i],LIGHT)
            margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; font(p.add_run(str(v)),size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t

def page_break(doc): doc.add_page_break()

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(.75); sec.left_margin=sec.right_margin=Inches(1)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in (("Heading 1",16,NAVY,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,"1F4D78",10,5)):
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

# Running header/footer
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(h.add_run("ATG PC AUDIT  |  Hướng dẫn sử dụng"),size=9,bold=True,color="6B7280")
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); f._p.append(fld)

# Cover: editorial manual pattern
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(36)
logo=ROOT/"assets"/"logo.png"
if logo.exists(): p.add_run().add_picture(str(logo),width=Inches(1.25))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20); font(p.add_run("ATG PC AUDIT"),size=30,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("HƯỚNG DẪN THAO TÁC SỬ DỤNG"),size=20,bold=True,color=BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(28); font(p.add_run("Kiểm tra máy tính • Quy hoạch IP • Đồng bộ và tổng hợp tài sản CNTT • Hỗ trợ trực tiếp Zalo 0904143113"),size=12,color="4B5563")
callout(doc,"Đối tượng sử dụng","Nhân viên thực hiện kiểm tra máy tính và quản trị viên phụ trách tổng hợp dữ liệu toàn công ty.",NAVY,"EAF2F8")
table(doc,["Thông tin","Nội dung"],[["Phiên bản tài liệu","1.0"],["Ứng dụng","ATG PC AUDIT"],["Phạm vi","Windows 10/11"],["Cập nhật","19/07/2026"]],[1.7,4.8])
body(doc,"Tài liệu này hướng dẫn thao tác từ khi nhận bộ chương trình, nhập hồ sơ, quét máy, gửi dữ liệu cho quản trị đến quản lý cơ sở dữ liệu và đồng bộ Google Sheet.")

page_break(doc); heading(doc,"1. Tổng quan và chuẩn bị",1)
heading(doc,"1.1. Thành phần bộ chương trình",2)
table(doc,["Thành phần","Mục đích"],[["ATG_PC_AUDIT.exe","Chương trình chính dùng để quét, gửi và tổng hợp dữ liệu."],["data\\config\\app_config.json","Cấu hình biểu mẫu, đồng bộ, sao lưu và các tùy chọn của ứng dụng."],["Kết quả kiểm tra","Thư mục tự động lưu CSV sau mỗi lần quét thành công."],["Biểu mẫu danh mục Excel","Xuất/nhập danh mục Loại máy tính, Phòng ban và Vị trí làm việc."],["ATG_PC_AUDIT_RECOVERY.exe","Công cụ quản trị dùng khôi phục mật khẩu khi cần."]],[2.25,4.25])
heading(doc,"1.2. Chuẩn bị trước khi chạy",2)
for x in ["Chép nguyên thư mục chương trình, không chỉ chép riêng file EXE.","Bảo đảm thư mục data nằm cùng cấp với file EXE để ứng dụng đọc app_config.json.","Kết nối Internet nếu cần gửi dữ liệu lên Google Sheet; khi mất mạng ứng dụng vẫn lưu CSV và hàng đợi cục bộ.","Đóng các bản ATG PC AUDIT đang chạy trước khi thay thế EXE mới."]:
    bullet(doc,x)
callout(doc,"Quan trọng","Không chỉnh sửa DEVICE_ID, device secret hoặc URL Google Apps Script nếu không có hướng dẫn của quản trị hệ thống.",RED,"FDECEC")
heading(doc,"1.3. Khởi động",2)
start_numbering(doc)
number(doc,"Nhấp đúp ATG_PC_AUDIT.exe.")
number(doc,"Nếu Windows hỏi quyền quản trị, chọn Yes. Quyền quản trị giúp ứng dụng đọc đầy đủ TPM, Secure Boot, bản quyền và phần cứng.")

page_break(doc); heading(doc,"2. Quy trình dành cho nhân viên",1)
heading(doc,"2.1. Nhập thông tin hồ sơ",2)
table(doc,["Trường","Cách nhập"],[["Loại máy tính","Chọn Bộ PC, Laptop hoặc Máy tính bảng từ danh sách."],["Người sử dụng","Nhập họ tên người đang sử dụng máy."],["Mã nhân viên","Nhập đúng mã nhân viên theo quy định công ty."],["Phòng ban","Chọn từ danh mục đã được quản trị cấu hình."],["Vị trí làm việc","Chọn tầng, phòng, văn phòng hoặc kho từ danh mục."],["Người thực hiện cập nhật","Nhập họ tên người kiểm tra máy."],["Ngày kiểm tra","Chọn ngày thực tế kiểm tra."],["Ghi chú","Ghi tình trạng đặc biệt: cần sao lưu, lỗi phần cứng, chuẩn bị cài lại Windows..."]],[1.65,4.85])
callout(doc,"Dấu * màu đỏ","Là trường bắt buộc. Ứng dụng không cho gửi dữ liệu nếu còn thiếu hoặc sai định dạng.",GOLD,"FFF7DD")
heading(doc,"2.2. Quét máy tính",2)
start_numbering(doc)
number(doc,"Kiểm tra lại thông tin hồ sơ.")
number(doc,"Bấm QUÉT MÁY TÍNH và chờ thanh tiến trình đạt 100%.")
number(doc,"Xem trạng thái “Đã hoàn thành kiểm tra máy tính”.")
number(doc,"Kiểm tra nhanh các tab Phần cứng, Windows, Mạng_IP và Bản quyền.")
number(doc,"Ứng dụng tự tạo một file CSV trong thư mục Kết quả kiểm tra và tự gửi lên máy chủ nếu có mạng.")

heading(doc,"2.3. Nội dung được thu thập",2)
table(doc,["Nhóm","Thông tin chính"],[["Máy tính","Tên máy, hãng, model, serial, UUID."],["CPU/RAM","Tên CPU, số nhân/luồng; tổng RAM và từng thanh RAM."],["Ổ cứng","Tổng dung lượng ổ đĩa vật lý chứa phân vùng Windows C:, model, serial và loại kết nối; không cộng các ổ dữ liệu khác."],["Card đồ họa","Tên GPU, driver, bộ nhớ đồ họa và trạng thái."],["Card mạng","Tên card, MAC, IPv4, gateway, DNS, DHCP và tốc độ kết nối."],["Windows 11","TPM, Secure Boot, UEFI, CPU, RAM, dung lượng và kết luận tương thích."],["Bản quyền","Trạng thái kích hoạt Windows và Microsoft Office."]],[1.45,5.05])

heading(doc,"3. Quy hoạch địa chỉ IP",1)
body(doc,"Tab Mạng_IP hiển thị card Ethernet/Wi-Fi vật lý và các card mạng khác. Dòng “Card mạng chính dùng cho quy hoạch IP” là giao diện đang được ưu tiên để lấy MAC và IP hiện tại.")
heading(doc,"3.1. Nhập kế hoạch IP",2)
table(doc,["Trường","Ý nghĩa"],[["VLAN dự kiến","VLAN sẽ cấp cho máy."],["IP dự kiến","Địa chỉ IP nội bộ dự kiến."],["Prefix","Độ dài subnet, ví dụ 24."],["Gateway dự kiến","Gateway của VLAN."],["DNS 1 / DNS 2","Máy chủ DNS nội bộ hoặc DNS theo chính sách."],["Tên switch / Cổng switch","Vị trí cắm mạng vật lý."],["Ổ cắm mạng","Mã ổ cắm tại bàn/phòng."],["Trạng thái triển khai","Chưa thực hiện, đang thực hiện hoặc hoàn thành."],["Ghi chú","Thông tin bổ sung phục vụ triển khai."]],[1.7,4.8])
heading(doc,"3.2. Kiểm tra trước khi áp dụng",2)
for x in ["Bấm KIỂM TRA IP DỰ KIẾN để phát hiện IP đang được sử dụng.","Chỉ bấm SAO CHÉP IP HIỆN TẠI SANG IP DỰ KIẾN khi muốn giữ nguyên địa chỉ đang dùng.","Ứng dụng chỉ ghi nhận kế hoạch; không tự thay đổi IP của Windows.","Dùng MAC chính để cấu hình DHCP reservation hoặc cố định IP trên hệ thống mạng."]:
    bullet(doc,x)

heading(doc,"4. Gửi và đồng bộ dữ liệu",1)
heading(doc,"4.1. Trạng thái gửi dữ liệu",2)
table(doc,["Thông báo","Ý nghĩa / xử lý"],[["Đã cập nhật","Dữ liệu đã được máy chủ tiếp nhận."],["Có dữ liệu mới","Có thay đổi và máy quản trị cần đồng bộ."],["Không thể kết nối Google Apps Script","Kiểm tra mạng hoặc URL API; CSV vẫn được lưu cục bộ."],["Thiết bị chờ phê duyệt","Chỉ xuất hiện nếu máy chủ vẫn dùng quy trình PENDING cũ; quản trị cần đặt ACTIVE hoặc cập nhật script."],["ROLE_DENIED","Vai trò thiết bị không được phép thực hiện tác vụ."],["INVALID_DEVICE_SECRET","Secret máy không khớp bản ghi trên THIET_BI; không tự xóa secret cục bộ."],["License AGGREGATE không hợp lệ","Chức năng Tổng hợp cần license AGGREGATE còn hạn."]],[2.4,4.1])
heading(doc,"4.2. Khi mất Internet",2)
body(doc,"Ứng dụng tiếp tục quét và tạo CSV. Bản ghi chưa gửi được sẽ nằm trong hàng đợi cục bộ. Khi có mạng, mở lại ứng dụng hoặc bấm GỬI LẠI DỮ LIỆU CHỜ để gửi tiếp.")
callout(doc,"Không quét lặp liên tục","Nếu đã có CSV và trạng thái chỉ lỗi mạng, ưu tiên gửi lại hàng đợi để tránh tạo nhiều lần kiểm tra không cần thiết.",GOLD,"FFF7DD")

heading(doc,"5. Chức năng Tổng hợp dành cho quản trị",1)
heading(doc,"5.1. Điều kiện mở",2)
for x in ["Thiết bị quản trị có role AGGREGATE và status ACTIVE trên sheet THIET_BI.","License có product_code ATG_PC_AUDIT, feature_code AGGREGATE, status ACTIVE và chưa hết hạn.","Đã thiết lập mật khẩu quản trị cục bộ."]:
    bullet(doc,x)
start_numbering(doc); number(doc,"Mở tab Tổng hợp.")
number(doc,"Bấm KÍCH HOẠT / KIỂM TRA LICENSE nếu trạng thái chưa hợp lệ.")
number(doc,"Nhập mật khẩu quản trị và bấm MỞ CHỨC NĂNG TỔNG HỢP.")
number(doc,"Bấm ĐỒNG BỘ NGAY để tải thay đổi từ Google Sheet.")
heading(doc,"5.2. Danh sách tổng hợp",2)
body(doc,"Bảng hiển thị mỗi máy đang hoạt động một dòng. Có thể tìm theo tên máy, serial, người dùng, mã nhân viên, phòng ban và các trường phần cứng.")
table(doc,["Thao tác","Cách dùng"],[["Xem chi tiết","Nhấp đúp vào dòng máy."],["Lọc dữ liệu","Nhập điều kiện và bấm ÁP DỤNG BỘ LỌC."],["Xóa khỏi danh sách hoạt động","Chọn dòng, bấm XÓA DÒNG ĐANG CHỌN và xác nhận. Lịch sử vẫn được giữ."],["Xuất báo cáo","Bấm XUẤT BÁO CÁO EXCEL và chọn nơi lưu."],["Import CSV","Chọn file → XEM TRƯỚC → XÁC NHẬN IMPORT."],["Lịch sử sử dụng","Mở chi tiết máy để xem các lần điều chuyển nhân viên và thay đổi phần cứng."]],[2.1,4.4])

page_break(doc); heading(doc,"6. Quản lý dữ liệu, sao lưu và phục hồi",1)
heading(doc,"6.1. Vị trí lưu dữ liệu",2)
body(doc,"Database, cấu hình, log, file xuất và backup được lưu trong thư mục dữ liệu đã chọn. Không nên đặt database trong thư mục tạm hoặc chỉ bên trong gói cài đặt có thể bị xóa.")
table(doc,["Nút","Chức năng"],[["MỞ THƯ MỤC DATABASE","Mở nơi chứa atg_pc_audit_master.db."],["MỞ THƯ MỤC CONFIG","Mở nơi chứa app_config.json."],["THAY ĐỔI THƯ MỤC DỮ LIỆU","Chuyển database và cấu hình sang vị trí mới."],["CHỌN THƯ MỤC BACKUP","Đổi nơi lưu bản sao lưu."],["SAO LƯU NGAY","Tạo file .atgbackup ngay lập tức."],["KHÔI PHỤC DỮ LIỆU","Chọn bản backup và phục hồi sau xác thực quản trị."],["KIỂM TRA DATABASE","Kiểm tra tính toàn vẹn cơ sở dữ liệu."]],[2.35,4.15])
heading(doc,"6.2. Chuyển sang máy quản trị mới",2)
start_numbering(doc)
number(doc,"Trên máy cũ, bấm SAO LƯU NGAY và chép file .atgbackup sang thiết bị lưu trữ an toàn.")
number(doc,"Chép nguyên bộ chương trình sang máy mới.")
number(doc,"Khởi động bằng thiết bị đã có license AGGREGATE hoặc đăng ký thiết bị mới.")
number(doc,"Chọn KHÔI PHỤC TỪ FILE BACKUP trong thiết lập dữ liệu, hoặc dùng nút KHÔI PHỤC DỮ LIỆU.")
number(doc,"Kiểm tra database rồi bấm ĐỒNG BỘ NGAY.")
callout(doc,"An toàn dữ liệu","Luôn sao lưu trước khi đổi đường dẫn, phục hồi, cập nhật phiên bản lớn hoặc chuyển máy.",RED,"FDECEC")

page_break(doc); heading(doc,"7. Quản lý biểu mẫu danh mục",1)
body(doc,"Biểu mẫu danh mục giúp toàn công ty dùng thống nhất Loại máy tính, Phòng ban và Vị trí làm việc. Các trường này hiển thị dạng danh sách thả xuống.")
heading(doc,"7.1. Xuất biểu mẫu",2)
start_numbering(doc)
number(doc,"Bấm nút CÀI ĐẶT/DANH MỤC cạnh khu vực nhập hồ sơ.")
number(doc,"Chọn XUẤT BIỂU MẪU và lưu file Excel.")
number(doc,"Mở Excel, bổ sung dữ liệu theo đúng từng cột; không đổi tên tiêu đề cột.")
heading(doc,"7.2. Nhập biểu mẫu",2)
start_numbering(doc)
number(doc,"Trong cửa sổ cài đặt danh mục, chọn NHẬP BIỂU MẪU.")
number(doc,"Chọn file Excel đã cập nhật và xác nhận.")
number(doc,"Đóng/mở lại danh sách hoặc ứng dụng để kiểm tra các lựa chọn mới.")
heading(doc,"7.3. Phân phối cấu hình",2)
body(doc,"Sau khi danh mục được lưu, chép file data\\config\\app_config.json cùng toàn bộ thư mục data sang máy người dùng. Cấu trúc bắt buộc:")
table(doc,["Cùng một thư mục","Bên trong"],[["ATG_PC_AUDIT.exe","File chạy"],["data","config\\app_config.json"],["Kết quả kiểm tra","Các file CSV tự động"]],[2.4,4.1])
callout(doc,"Nếu máy mới không tải danh mục","Kiểm tra file có đúng đường dẫn data\\config\\app_config.json, ứng dụng đã đóng khi chép file và JSON không bị lỗi cú pháp.",GOLD,"FFF7DD")

page_break(doc); heading(doc,"8. Giải thích các tab kết quả",1)
table(doc,["Tab","Nội dung"],[["Tổng quan","Thông tin hồ sơ, máy, trạng thái Windows 11, bản quyền, MAC và IP chính."],["Phần cứng","Máy tính, CPU, RAM, BIOS/Mainboard, ổ đĩa vật lý, GPU và card mạng."],["Windows","Thông tin Windows hiện tại và đánh giá điều kiện Windows 11."],["Mạng_IP","Card mạng vật lý, MAC/IP và kế hoạch IP nội bộ."],["Bản quyền","Kết quả kích hoạt Windows và Office."],["Xuất dữ liệu","CSV dự phòng, trạng thái gửi máy chủ và gửi lại hàng đợi."],["Tổng hợp","Đồng bộ toàn công ty, import/export, database, backup và báo cáo; chỉ dành cho quản trị."]],[1.4,5.1])
heading(doc,"8.1. Đọc thông tin ổ cứng",2)
body(doc,"Cột Dung lượng ổ cài HĐH (GB) là tổng dung lượng của ổ đĩa vật lý chứa phân vùng Windows C:. Ví dụ phân vùng C: chỉ có 240 GB nhưng nằm trên ổ vật lý 476,92 GB thì báo cáo hiển thị khoảng 476,92 GB. Các ổ dữ liệu vật lý khác không được cộng vào giá trị này.")
heading(doc,"8.2. Đọc lịch sử sử dụng máy",2)
body(doc,"Ứng dụng đối chiếu UUID và serial để nhận dạng cùng một máy. Khi người dùng, mã nhân viên, phòng ban hoặc vị trí thay đổi, hệ thống tạo lần sử dụng mới. Nếu CPU, RAM, ổ cứng, GPU, serial hoặc model thay đổi, ghi chú lịch sử nêu rõ giá trị cũ và mới.")

page_break(doc); heading(doc,"9. Xử lý sự cố thường gặp",1)
table(doc,["Hiện tượng","Cách xử lý nhanh"],[["Smart App Control chặn EXE","Dùng bản EXE đã ký số nếu triển khai rộng. Trên máy thử nghiệm, quản trị có thể tắt Smart App Control theo chính sách công ty; chạy Administrator không vượt qua được cơ chế này."],["Quét thiếu TPM/Secure Boot/bản quyền","Đóng ứng dụng và chạy lại bằng quyền Administrator."],["Không thấy GPU","Cập nhật EXE mới, quét lại và kiểm tra tab Card đồ họa/GPU."],["ID ổ cứng Không xác định","Dùng EXE mới và quét lại. Ứng dụng sẽ dùng Disk ID ổ vật lý đầu tiên nếu không xác định được ổ chứa C:."],["Google Sheet có dữ liệu nhưng Tổng hợp thiếu máy","Kiểm tra role/status thiết bị, license AGGREGATE, URL deployment mới nhất; sau đó khóa/mở lại Tổng hợp và đồng bộ."],["Device secret không khớp","Không xóa secret cục bộ. Xóa đúng bản ghi thiết bị cũ trên THIET_BI rồi đăng ký lại theo quy trình quản trị."],["Không tải được app_config.json","Chép nguyên thư mục data cạnh EXE, kiểm tra đường dẫn và khởi động lại ứng dụng."],["Mất mật khẩu quản trị","Dùng ATG_PC_AUDIT_RECOVERY.exe đặt mật khẩu mới; công cụ tạo backup trước khi thay đổi."]],[2.4,4.1])

page_break(doc); heading(doc,"10. Quy trình chuẩn đề xuất",1)
heading(doc,"10.1. Nhân viên/IT kiểm tra máy",2)
for x in ["Nhận danh mục và app_config.json mới nhất.","Mở ứng dụng bằng quyền quản trị.","Chọn đúng Loại máy, Phòng ban, Vị trí; nhập người dùng và mã nhân viên.","Bấm QUÉT MÁY TÍNH; kiểm tra CPU, RAM, dung lượng ổ vật lý cài Windows, GPU, MAC/IP và Windows 11.","Nhập kế hoạch IP nếu có.","Xác nhận CSV đã tạo và trạng thái gửi dữ liệu thành công; nếu mất mạng, giữ CSV và gửi lại sau."]:
    bullet(doc,x)
heading(doc,"10.2. Quản trị viên",2)
for x in ["Duy trì danh mục chuẩn và phân phối app_config.json.","Kiểm tra thiết bị AGGREGATE, license và trạng thái API.","Mở Tổng hợp, bấm ĐỒNG BỘ NGAY và đối chiếu số lượng máy.","Rà soát xung đột, lịch sử điều chuyển và thay đổi phần cứng.","Xuất báo cáo định kỳ; sao lưu database trước và sau thay đổi quan trọng.","Theo dõi thư mục backup và kiểm tra khả năng phục hồi định kỳ."]:
    bullet(doc,x)
callout(doc,"Hoàn thành kiểm tra khi","Thông tin hồ sơ đầy đủ • quét đạt 100% • CSV đã tạo • dữ liệu đã gửi hoặc nằm an toàn trong hàng đợi • quản trị nhìn thấy máy sau đồng bộ.",GREEN,"E8F5E9")

doc.core_properties.title="Hướng dẫn sử dụng ATG PC AUDIT"
doc.core_properties.subject="Hướng dẫn thao tác cho nhân viên và quản trị viên"
doc.core_properties.author="ATG PC AUDIT"
doc.save(OUT)
print(OUT)
